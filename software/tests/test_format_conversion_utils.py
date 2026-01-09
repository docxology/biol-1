"""Comprehensive tests for format conversion utilities."""

from pathlib import Path
import pytest

from src.format_conversion.utils import (
    convert_docx_to_markdown,
    convert_markdown_to_docx,
    convert_markdown_to_html,
    convert_pdf_to_text,
    get_conversion_path,
    get_file_extension,
    ensure_output_directory,
    get_output_path,
    convert_markdown_to_pdf,
    convert_html_to_pdf,
    convert_text_to_pdf,
    convert_text_to_html,
    convert_audio_to_text,
)


class TestGetFileExtension:
    """Tests for get_file_extension function."""

    def test_get_file_extension_md(self):
        """Test getting markdown extension."""
        result = get_file_extension(Path("test.md"))
        assert result == "md"

    def test_get_file_extension_pdf(self):
        """Test getting PDF extension."""
        result = get_file_extension(Path("test.pdf"))
        assert result == "pdf"

    def test_get_file_extension_uppercase(self):
        """Test getting uppercase extension."""
        result = get_file_extension(Path("test.PDF"))
        assert result == "pdf"


class TestEnsureOutputDirectory:
    """Tests for ensure_output_directory function."""

    def test_ensure_output_directory_creates(self, temp_dir):
        """Test that directory is created."""
        new_dir = temp_dir / "newdir" / "subdir"
        ensure_output_directory(new_dir)
        assert new_dir.parent.exists()

    def test_ensure_output_directory_existing(self, temp_dir):
        """Test with existing directory."""
        ensure_output_directory(temp_dir / "file.txt")
        # Should not raise


class TestGetOutputPath:
    """Tests for get_output_path function."""

    def test_get_output_path_same_dir(self, temp_dir):
        """Test output path in same directory."""
        input_path = temp_dir / "test.md"
        result = get_output_path(input_path, "pdf")
        assert result.suffix == ".pdf"
        assert result.parent == temp_dir

    def test_get_output_path_different_dir(self, temp_dir):
        """Test output path in different directory."""
        input_path = temp_dir / "test.md"
        output_dir = temp_dir / "output"
        result = get_output_path(input_path, "html", output_dir)
        assert result.suffix == ".html"
        assert result.parent == output_dir


class TestConvertMarkdownToHtml:
    """Tests for convert_markdown_to_html function."""

    def test_convert_markdown_to_html_heading(self, temp_dir):
        """Test converting markdown heading to HTML."""
        md_file = temp_dir / "test.md"
        md_file.write_text("# Hello World\n\nParagraph text.", encoding="utf-8")
        html_file = temp_dir / "test.html"

        convert_markdown_to_html(md_file, html_file)

        content = html_file.read_text()
        assert html_file.exists()

    def test_convert_markdown_to_html_list(self, temp_dir):
        """Test converting markdown list to HTML."""
        md_file = temp_dir / "test.md"
        md_file.write_text("- Item 1\n- Item 2\n- Item 3", encoding="utf-8")
        html_file = temp_dir / "test.html"

        convert_markdown_to_html(md_file, html_file)

        assert html_file.exists()


class TestConvertMarkdownToPdf:
    """Tests for convert_markdown_to_pdf function."""

    def test_convert_markdown_to_pdf_basic(self, temp_dir):
        """Test converting markdown to PDF."""
        md_file = temp_dir / "test.md"
        md_file.write_text("# Title\n\nContent", encoding="utf-8")
        pdf_file = temp_dir / "test.pdf"

        convert_markdown_to_pdf(md_file, pdf_file)

        assert pdf_file.exists()


class TestConvertHtmlToPdf:
    """Tests for convert_html_to_pdf function."""

    def test_convert_html_to_pdf_basic(self, temp_dir):
        """Test converting HTML to PDF."""
        html_file = temp_dir / "test.html"
        html_file.write_text("<html><body><h1>Title</h1></body></html>", encoding="utf-8")
        pdf_file = temp_dir / "test.pdf"

        convert_html_to_pdf(html_file, pdf_file)

        assert pdf_file.exists()


class TestConvertTextToPdf:
    """Tests for convert_text_to_pdf function."""

    def test_convert_text_to_pdf_basic(self, temp_dir):
        """Test converting text to PDF."""
        txt_file = temp_dir / "test.txt"
        txt_file.write_text("Plain text content.", encoding="utf-8")
        pdf_file = temp_dir / "test.pdf"

        convert_text_to_pdf(txt_file, pdf_file)

        assert pdf_file.exists()


class TestConvertTextToHtml:
    """Tests for convert_text_to_html function."""

    def test_convert_text_to_html_basic(self, temp_dir):
        """Test converting text to HTML."""
        txt_file = temp_dir / "test.txt"
        txt_file.write_text("Line 1\nLine 2", encoding="utf-8")
        html_file = temp_dir / "test.html"

        convert_text_to_html(txt_file, html_file)

        assert html_file.exists()
        content = html_file.read_text()
        assert "<html>" in content or "<p>" in content


class TestConvertMarkdownToDocx:
    """Tests for convert_markdown_to_docx function."""

    def test_convert_markdown_to_docx_basic(self, temp_dir):
        """Test converting markdown to DOCX."""
        md_file = temp_dir / "test.md"
        md_file.write_text("# Title\n\nParagraph content.", encoding="utf-8")
        docx_file = temp_dir / "test.docx"

        convert_markdown_to_docx(md_file, docx_file)

        assert docx_file.exists()

    def test_convert_markdown_to_docx_with_list(self, temp_dir):
        """Test converting markdown list to DOCX."""
        md_file = temp_dir / "test.md"
        md_file.write_text("## Items\n\n- One\n- Two\n- Three", encoding="utf-8")
        docx_file = temp_dir / "test.docx"

        convert_markdown_to_docx(md_file, docx_file)

        assert docx_file.exists()


class TestConvertDocxToMarkdown:
    """Tests for convert_docx_to_markdown function."""

    def test_convert_docx_to_markdown_basic(self, temp_dir):
        """Test converting DOCX to markdown."""
        from docx import Document

        # Create a simple DOCX file
        docx_file = temp_dir / "test.docx"
        doc = Document()
        doc.add_heading("Test Title", level=1)
        doc.add_paragraph("Test paragraph content.")
        doc.save(str(docx_file))

        result = convert_docx_to_markdown(docx_file)

        assert "Test Title" in result


class TestConvertPdfToText:
    """Tests for convert_pdf_to_text function."""

    def test_convert_pdf_to_text_basic(self, temp_dir):
        """Test converting PDF to text."""
        # Create a simple PDF using markdown_to_pdf
        from src.markdown_to_pdf.main import render_markdown_to_pdf

        md_file = temp_dir / "test.md"
        md_file.write_text("# PDF Test\n\nThis is test content for PDF.", encoding="utf-8")
        pdf_file = temp_dir / "test.pdf"
        render_markdown_to_pdf(str(md_file), str(pdf_file))

        txt_file = temp_dir / "output.txt"
        convert_pdf_to_text(pdf_file, txt_file)

        content = txt_file.read_text()
        assert len(content) > 0


class TestGetConversionPath:
    """Tests for get_conversion_path function."""

    def test_get_conversion_path_md_to_pdf(self):
        """Test generating path for markdown to PDF."""
        result = get_conversion_path("/path/to/file.md", "pdf")
        assert result == "/path/to/file.pdf"

    def test_get_conversion_path_md_to_html(self):
        """Test generating path for markdown to HTML."""
        result = get_conversion_path("/path/to/file.md", "html")
        assert result == "/path/to/file.html"
