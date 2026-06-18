"""Utility functions for format conversion."""

from html.parser import HTMLParser
from pathlib import Path
from typing import Optional

from ..markdown_to_pdf.main import render_markdown_to_pdf
from ..markdown_to_pdf.utils import markdown_to_html
from src.shared.file_utils import read_markdown_file


def get_file_extension(file_path: Path) -> str:
    """Get file extension (lowercase, without dot).

    Args:
        file_path: Path to file

    Returns:
        File extension without dot
    """
    return file_path.suffix.lower().lstrip(".")


def get_output_path(
    input_path: Path, output_format: str, output_dir: Optional[Path] = None
) -> Path:
    """Get output file path from input path and format.

    Args:
        input_path: Path to input file
        output_format: Target format (without dot)
        output_dir: Optional output directory (if None, uses input directory)

    Returns:
        Path to output file
    """
    if output_dir is None:
        output_dir = input_path.parent

    output_filename = input_path.stem + "." + output_format
    return output_dir / output_filename


def convert_markdown_to_pdf(input_path: Path, output_path: Path) -> None:
    """Convert Markdown file to PDF."""
    render_markdown_to_pdf(str(input_path), str(output_path))


def convert_markdown_to_html(input_path: Path, output_path: Path) -> None:
    """Convert Markdown file to HTML.

    Args:
        input_path: Path to input Markdown file
        output_path: Path to output HTML file
    """
    # imports moved to top-level

    markdown_content = read_markdown_file(input_path)
    html_content = markdown_to_html(markdown_content)

    # Add basic HTML structure
    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{input_path.stem}</title>
</head>
<body>
{html_content}
</body>
</html>"""

    output_path.write_text(full_html, encoding="utf-8")


def convert_html_to_pdf(input_path: Path, output_path: Path) -> None:
    """Convert HTML file to PDF.

    Args:
        input_path: Path to input HTML file
        output_path: Path to output PDF file
    """
    from weasyprint import HTML

    html_content = input_path.read_text(encoding="utf-8")
    HTML(string=html_content).write_pdf(output_path)


def convert_text_to_pdf(input_path: Path, output_path: Path) -> None:
    """Convert text file to PDF.

    Args:
        input_path: Path to input text file
        output_path: Path to output PDF file
    """
    from weasyprint import HTML

    text_content = input_path.read_text(encoding="utf-8")
    # Escape HTML and wrap in pre tag
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        pre {{ font-family: monospace; white-space: pre-wrap; }}
    </style>
</head>
<body>
    <pre>{text_content}</pre>
</body>
</html>"""
    HTML(string=html_content).write_pdf(output_path)


def convert_text_to_html(input_path: Path, output_path: Path) -> None:
    """Convert text file to HTML.

    Args:
        input_path: Path to input text file
        output_path: Path to output HTML file
    """
    text_content = input_path.read_text(encoding="utf-8")
    # Escape HTML and wrap in pre tag
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{input_path.stem}</title>
</head>
<body>
    <pre>{text_content}</pre>
</body>
</html>"""
    output_path.write_text(html_content, encoding="utf-8")


def convert_markdown_to_docx(input_path: Path, output_path: Path) -> None:
    """Convert a Markdown file to DOCX.

    Renders headings, paragraphs, ordered/unordered lists (including
    nested lists and tight lists with no blank lines between items),
    inline emphasis (``**bold**``, ``*italic*``, ``` `code` ```),
    blockquotes, fenced code blocks, and simple tables.

    Args:
        input_path: Path to input Markdown file
        output_path: Path to output DOCX file
    """
    from docx import Document

    markdown_content = read_markdown_file(input_path)
    html_content = markdown_to_html(markdown_content)

    doc = Document()
    walker = _MarkdownHtmlToDocx(doc)
    walker.feed(html_content)
    walker.close()
    walker.finalize()
    doc.save(str(output_path))


class _MarkdownHtmlToDocx(HTMLParser):
    """Walk markdown-derived HTML and emit DOCX paragraphs and runs.

    Treats ``p``, ``h1``-``h6``, ``li``, ``blockquote``, and ``pre`` as
    block boundaries that flush their accumulated runs into a paragraph.
    Maintains a stack for nested ordered/unordered lists so tight lists
    (no blank lines between items) render correctly.
    """

    BLOCK_TAGS = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "blockquote", "pre"}
    EMPHASIS_TAGS = {"strong", "b", "em", "i", "code", "u"}

    def __init__(self, doc) -> None:
        super().__init__(convert_charrefs=True)
        self._doc = doc
        self._block_stack: list[str] = []
        self._list_stack: list[dict] = []
        self._fmt: dict[str, int] = {"bold": 0, "italic": 0, "code": 0, "underline": 0}
        self._runs: list[tuple[str, dict]] = []
        self._in_table = False
        self._table_rows: list[list[str]] = []
        self._current_row: list[str] = []
        self._current_cell: list[str] = []

    def finalize(self) -> None:
        """Flush any trailing block content (e.g. unclosed paragraph)."""
        if self._runs and self._block_stack:
            self._flush_block(self._block_stack[-1])

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in ("ol", "ul"):
            self._list_stack.append({"type": tag, "index": 0})
            return
        if tag == "li":
            if self._list_stack:
                self._list_stack[-1]["index"] += 1
            self._block_stack.append("li")
            return
        if tag == "p" and self._block_stack and self._block_stack[-1] == "li":
            return
        if tag in self.BLOCK_TAGS:
            self._block_stack.append(tag)
            return
        if tag in self.EMPHASIS_TAGS:
            self._fmt[self._fmt_key(tag)] += 1
            return
        if tag == "br":
            self._runs.append(("\n", self._snapshot_fmt()))
            return
        if tag == "table":
            self._in_table = True
            self._table_rows = []
            return
        if tag == "tr" and self._in_table:
            self._current_row = []
            return
        if tag in ("td", "th") and self._in_table:
            self._current_cell = []
            return

    def handle_endtag(self, tag: str) -> None:
        if tag in ("ol", "ul"):
            if self._list_stack:
                self._list_stack.pop()
            return
        if tag == "p" and self._block_stack and self._block_stack[-1] == "li":
            return
        if tag in self.BLOCK_TAGS:
            self._flush_block(tag)
            if self._block_stack and self._block_stack[-1] == tag:
                self._block_stack.pop()
            return
        if tag in self.EMPHASIS_TAGS:
            key = self._fmt_key(tag)
            if self._fmt[key] > 0:
                self._fmt[key] -= 1
            return
        if tag in ("td", "th") and self._in_table:
            self._current_row.append("".join(self._current_cell).strip())
            self._current_cell = []
            return
        if tag == "tr" and self._in_table:
            self._table_rows.append(self._current_row)
            self._current_row = []
            return
        if tag == "table" and self._in_table:
            self._emit_table(self._table_rows)
            self._in_table = False
            self._table_rows = []
            return

    def handle_data(self, data: str) -> None:
        if self._in_table:
            self._current_cell.append(data)
            return
        if not self._block_stack:
            return
        self._runs.append((data, self._snapshot_fmt()))

    @staticmethod
    def _fmt_key(tag: str) -> str:
        if tag in ("strong", "b"):
            return "bold"
        if tag in ("em", "i"):
            return "italic"
        if tag == "u":
            return "underline"
        return "code"

    def _snapshot_fmt(self) -> dict:
        return {k: v > 0 for k, v in self._fmt.items()}

    def _list_prefix(self) -> str:
        if not self._list_stack:
            return ""
        ctx = self._list_stack[-1]
        return f"{ctx['index']}. " if ctx["type"] == "ol" else "• "

    def _flush_block(self, tag: str) -> None:
        runs = self._runs
        self._runs = []
        if not any(text.strip() for text, _ in runs):
            return

        if len(tag) == 2 and tag[0] == "h" and tag[1].isdigit():
            level = min(int(tag[1]), 6)
            paragraph = self._doc.add_heading("", level=level)
        elif tag == "blockquote":
            try:
                paragraph = self._doc.add_paragraph(style="Quote")
            except KeyError:
                paragraph = self._doc.add_paragraph()
        else:
            paragraph = self._doc.add_paragraph()

        if tag == "li":
            paragraph.add_run(self._list_prefix())

        for text, fmt in runs:
            if not text:
                continue
            run = paragraph.add_run(text)
            if fmt.get("bold"):
                run.bold = True
            if fmt.get("italic"):
                run.italic = True
            if fmt.get("underline"):
                run.underline = True
            if fmt.get("code"):
                run.font.name = "Courier New"

    def _emit_table(self, rows: list[list[str]]) -> None:
        if not rows:
            return
        n_cols = max(len(r) for r in rows)
        table = self._doc.add_table(rows=len(rows), cols=n_cols)
        for r_idx, row in enumerate(rows):
            for c_idx in range(n_cols):
                table.cell(r_idx, c_idx).text = row[c_idx] if c_idx < len(row) else ""


def convert_docx_to_markdown(input_path: Path) -> str:
    """Convert DOCX file to Markdown format.

    Args:
        input_path: Path to input DOCX file

    Returns:
        Markdown content as string
    """
    from docx import Document

    doc = Document(str(input_path))
    markdown_lines = []

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            markdown_lines.append("")
            continue

        # Check if paragraph is a heading
        style = paragraph.style
        style_name = style.name.lower() if style is not None and style.name else ""
        if "heading" in style_name:
            level = 1
            if "heading 2" in style_name or "subtitle" in style_name:
                level = 2
            elif "heading 3" in style_name:
                level = 3
            elif "heading 4" in style_name:
                level = 4
            elif "heading 5" in style_name:
                level = 5
            elif "heading 6" in style_name:
                level = 6

            text = _extract_formatted_text(paragraph)
            markdown_lines.append(f"{'#' * level} {text}")
        else:
            # Regular paragraph - extract formatted text
            text = _extract_formatted_text(paragraph)
            if text.strip():
                markdown_lines.append(text)

        # Add blank line after paragraph
        markdown_lines.append("")

    # Process tables
    for table in doc.tables:
        markdown_lines.append("")
        # Extract table header
        if table.rows:
            header_row = table.rows[0]
            header_cells = [
                _extract_formatted_text(cell) for cell in header_row.cells
            ]
            markdown_lines.append("| " + " | ".join(header_cells) + " |")
            markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

            # Extract data rows
            for row in table.rows[1:]:
                cells = [_extract_formatted_text(cell) for cell in row.cells]
                markdown_lines.append("| " + " | ".join(cells) + " |")
        markdown_lines.append("")

    # Remove trailing blank lines
    while markdown_lines and not markdown_lines[-1].strip():
        markdown_lines.pop()

    return "\n".join(markdown_lines)


def _extract_formatted_text(paragraph) -> str:
    """Extract text from paragraph with formatting preserved as Markdown.

    Args:
        paragraph: docx paragraph object

    Returns:
        Formatted text string
    """
    text_parts = []

    for run in paragraph.runs:
        text = run.text
        if not text:
            continue

        # Apply formatting
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        if run.underline:
            text = f"<u>{text}</u>"

        text_parts.append(text)

    return "".join(text_parts)


def convert_pdf_to_text(input_path: Path, output_path: Path) -> None:
    """Convert PDF file to text.

    Args:
        input_path: Path to input PDF file
        output_path: Path to output text file
    """
    from pypdf import PdfReader

    reader = PdfReader(str(input_path))
    text_content = []

    for page in reader.pages:
        text_content.append(page.extract_text())

    full_text = "\n\n".join(text_content)
    output_path.write_text(full_text, encoding="utf-8")


def convert_audio_to_text(input_path: Path, output_path: Path) -> None:
    """Convert audio file to text using speech-to-text.

    Args:
        input_path: Path to input audio file
        output_path: Path to output text file
    """
    from ..speech_to_text.main import transcribe_audio

    transcribe_audio(str(input_path), str(output_path))


def get_conversion_path(input_path: str, output_format: str) -> str:
    """Generate output path for file conversion.

    Thin string-typed wrapper around get_output_path for callers
    that operate on string paths rather than Path objects.

    Args:
        input_path: Path to input file
        output_format: Target format (without dot)

    Returns:
        Output file path with new extension
    """
    return str(get_output_path(Path(input_path), output_format))
